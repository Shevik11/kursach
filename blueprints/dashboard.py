import os
import logging
from flask import Blueprint, render_template, request, session, redirect, url_for, flash
from extensions import db, mail
from firebase_admin import firestore
from werkzeug.utils import secure_filename
from docx import Document
from datetime import datetime
from flask_mail import Message
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from dotenv import load_dotenv

# Завантаження змінних середовища
load_dotenv()

# Налаштування логування
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ініціалізація токенів
s = URLSafeTimedSerializer(os.getenv('SECRET_KEY'))

# Налаштування завантаження файлів
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

# Ініціалізація Blueprint
dashboard_blueprint = Blueprint('dashboard', __name__)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_path(username, filename):
    return os.path.join(UPLOAD_FOLDER, username, secure_filename(filename))

@dashboard_blueprint.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))

    user_id = session['user_id']
    user_ref = db.collection('users').document(user_id)
    user_data = user_ref.get().to_dict()

    if not user_data:
        flash("User data not found.", 'error')
        return redirect(url_for('auth.login'))

    username = user_data.get('username', user_id)
    user_folder = os.path.join(UPLOAD_FOLDER, username)
    os.makedirs(user_folder, exist_ok=True)

    local_files = [
        {'name': filename, 'url': f"/uploads/{username}/{filename}"}
        for filename in os.listdir(user_folder)
        if allowed_file(filename)
    ]
    local_files.sort(key=lambda x: x['name'])

    if request.method == 'POST':
        file = request.files['file']
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(user_folder, filename)

            if filename not in user_data.get('files', []):
                try:
                    user_ref.update({'files': firestore.ArrayUnion([filename])})
                    file.save(file_path)
                    flash(f"File '{filename}' uploaded successfully!", 'success')
                    local_files.append({'name': filename, 'url': f"/uploads/{username}/{filename}"})
                except Exception as e:
                    logger.error(f"Error uploading file: {str(e)}")
                    flash(f"Error uploading file: {str(e)}", 'error')
            else:
                flash(f"File '{filename}' already exists.", 'info')

    return render_template('index.html', files=local_files)


@dashboard_blueprint.route('/delete_file/<filename>', methods=['POST'])
def delete_file(filename):
    try:
        user_id = session['user_id']
        user_ref = db.collection('users').document(user_id)
        user_data = user_ref.get().to_dict()

        username = user_data.get('username', user_id)
        file_path = get_file_path(username, filename)

        if os.path.exists(file_path):
            os.remove(file_path)
            user_ref.update({'files': firestore.ArrayRemove([filename])})
            flash(f"File '{filename}' deleted successfully.", 'success')
        else:
            flash(f"File '{filename}' does not exist.", 'error')

        return redirect(url_for('dashboard.dashboard'))
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        flash(f"Error deleting file: {str(e)}", 'error')
        return redirect(url_for('dashboard.dashboard'))


@dashboard_blueprint.route('/view_file/<filename>', methods=['GET'])
def view_file(filename):
    try:
        user_id = session['user_id']
        username = db.collection('users').document(user_id).get().to_dict().get('username', user_id)
        file_path = os.path.join(UPLOAD_FOLDER, username, filename)

        # Перевірка, чи існує файл
        if not os.path.exists(file_path):
            flash(f"File '{filename}' does not exist.", 'error')
            return redirect(url_for('dashboard.dashboard'))

        content = ""
        # Зчитування вмісту файлу
        if filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
        else:
            flash("Unsupported file type.", 'error')
            return redirect(url_for('dashboard.dashboard'))

        return render_template('view_file.html', filename=filename, content=content)
    except Exception as e:
        flash(f"Error reading file: {str(e)}", 'error')
        return redirect(url_for('dashboard.dashboard'))


@dashboard_blueprint.route('/edit_file/<filename>', methods=['GET', 'POST'])
def edit_file(filename):
    try:
        user_id = session['user_id']
        user_ref = db.collection('users').document(user_id)
        user_data = user_ref.get().to_dict()

        username = user_data.get('username', user_id)
        file_path = os.path.join(UPLOAD_FOLDER, username, filename)

        if not os.path.exists(file_path):
            flash(f"Файл '{filename}' не існує.", 'error')
            return redirect(url_for('dashboard.dashboard'))

        # Завантаження поточного вмісту файлу для відображення
        current_content = ""
        if filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                current_content = f.read()
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            current_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        if request.method == 'POST':
            # Отримання вмісту та опису з форми
            new_content = request.form.get('content')
            print(f"New content: {new_content}")
            if not new_content:
                flash("Не вказано новий вміст для файлу.", 'error')
                return redirect(url_for('dashboard.dashboard'))
            session['new_content'] = new_content
            description = request.form.get('description', '')

            # Збереження даних у сесії
            session['description'] = description

            # Створення токену для підтвердження змін
            token_data = {
                'filename': filename,
                'user_id': user_id,
                'new_content': new_content,  # Зберігаємо текст без змін
                'description': description
            }
            token = s.dumps(token_data, salt='confirm-change')
            confirm_url = url_for('dashboard.confirm_change', token=token, _external=True, filename=filename, content=new_content)

            # Відправлення підтвердження на пошту
            msg = Message('Підтвердження змін у файлі', recipients=[user_data['email']])
            msg.body = f"Щоб підтвердити зміни до файлу {filename}, перейдіть за посиланням: {confirm_url}"
            try:
                mail.send(msg)
                flash("Зміни збережено. Перевірте свою пошту для підтвердження.", 'success')
            except Exception as e:
                logger.error(f"Не вдалося надіслати лист: {str(e)}")
                flash("Не вдалося надіслати лист для підтвердження. Спробуйте ще раз.", 'error')

            return redirect(url_for('dashboard.dashboard'))

        return render_template('edit_file.html', filename=filename, content=current_content)
    except Exception as e:
        logger.error(f"Помилка під час редагування файлу: {str(e)}")
        flash(f"Помилка під час редагування файлу: {str(e)}", 'error')
        return redirect(url_for('dashboard.dashboard'))

@dashboard_blueprint.route('/confirm_change/<token>/<filename>', methods=['GET'])
def confirm_change(token, filename):
    try:
        user_id = session.get('user_id')
        if not user_id:
            flash("Сеанс завершено. Будь ласка, увійдіть знову.", 'error')
            return redirect(url_for('auth.login'))

        data = s.loads(token, salt='confirm-change', max_age=3600)

        # Отримання вмісту з сесії
        new_content = session.get('new_content')
        description = session.get('description')

        if not new_content:
            logger.error("Не знайдено новий вміст для файлу у сесії.")
            flash("Не знайдено новий вміст для файлу.", 'error')
            return redirect(url_for('dashboard.dashboard'))

        logger.info(f"Знайдено новий вміст для файлу у сесії: {new_content}")

        # Отримання шляху до файлу
        username = db.collection('users').document(user_id).get().to_dict().get('username', user_id)
        file_path = os.path.join(UPLOAD_FOLDER, username, filename)

        # Оновлення файлу
        if filename.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(new_content)
        elif filename.endswith('.docx'):
            doc = Document()
            for line in new_content.split('\n'):
                doc.add_paragraph(line)
            doc.save(file_path)

        # Оновлення бази даних
        timestamp = datetime.utcnow().isoformat()
        user_ref = db.collection('users').document(user_id)
        user_ref.update({
            'file_changes': firestore.ArrayUnion([{
                'filename': filename,
                'description': description,
                'timestamp': timestamp,
                'modified_by': username
            }])
        })

        flash("Зміни підтверджено.", 'success')
        return redirect(url_for('dashboard.dashboard'))
    except SignatureExpired:
        flash("Посилання для підтвердження зміни втратило актуальність.", 'error')
    except BadSignature:
        flash("Недійсне посилання для підтвердження.", 'error')
    except Exception as e:
        logger.error(f"Помилка підтвердження змін: {str(e)}")
        flash(f"Помилка підтвердження змін: {str(e)}", 'error')

    return redirect(url_for('dashboard.dashboard'))
