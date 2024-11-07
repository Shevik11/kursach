import os
from flask import flash, render_template, request, redirect, url_for
import docx  # Додаємо бібліотеку для роботи з docx

def read_file_content(file_path):
    """Функція для читання вмісту різних типів файлів"""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='cp1251') as f:
                    return f.read()
            except UnicodeDecodeError:
                return "Помилка кодування файлу"
        except Exception as e:
            return f"Помилка при читанні файлу: {str(e)}"

    elif file_extension == '.docx':
        try:
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            return f"Помилка читання DOCX файлу: {str(e)}"

    else:
        return "Цей тип файлу не підтримується для редагування"


def save_file_content(file_path, content):
    """Функція для збереження вмісту файлів"""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.txt':
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            flash(f"Помилка збереження TXT файлу: {str(e)}", 'danger')
            return False

    elif file_extension == '.docx':
        try:
            doc = docx.Document()
            for paragraph in content.split('\n'):
                doc.add_paragraph(paragraph)
            doc.save(file_path)
            return True
        except Exception as e:
            flash(f"Помилка збереження DOCX файлу: {str(e)}", 'danger')
            return False

    return False
